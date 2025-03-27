import fitz
import aspose.words as aw
# import streamlit as st
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml import OxmlElement
from docx2pdf import convert
import xml.etree.ElementTree as ET
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sentence_transformers import SentenceTransformer
import torch
import io
import subprocess
import platform

# from dotenv import load_dotenv
from html.parser import HTMLParser

import os
import time
import tempfile
import re
import fitz  # PyMuPDF
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_BREAK
from bs4 import BeautifulSoup
from bs4.element import NavigableString
import difflib
import mimetypes

def allowed_file(filename):
    return filename.lower().endswith('.docx')

# def is_docx_empty(file_path):
#     try:
#         doc = Document(file_path)
#         return not any(p.text.strip() for p in doc.paragraphs)
#     except Exception:
#         return True
    
# def is_valid_docx(file):
#     """Check if the uploaded file is a valid DOCX file."""
#     mime_type, _ = mimetypes.guess_type(file)
#     print(">>>>>>>>>>>>>> mime_type", mime_type)
#     return mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def docx_to_plain_text1(temp_file_path):
    """
    Convert a .docx file to plain text format, maintaining the sequence of elements
    (paragraphs, tables, images) and tracking line and page numbers.

    Args:
        file: A file-like object containing the .docx file

    Returns:
        list: List of dictionaries containing line number, page number, and text content
    """
    # with tempfile.TemporaryDirectory() as temp_dir:
    #     # Save uploaded file to temp directory
    #     temp_file_path = os.path.join(temp_dir, "temp.docx")
    #     with open(temp_file_path, "wb") as f:
    #         f.write(file.read())

    # Load the document
    document = Document(temp_file_path)

    # Initialize trackers
    plain_text = []
    line_number = 1
    page_number = 1
    chars_on_page = 0
    chars_per_page = 3500  # Approximate characters per page

    # Function to add content and track page numbers
    def add_content(text, content_type="text"):
        nonlocal line_number, page_number, chars_on_page

        # Calculate page breaks based on character count
        chars_on_page += len(text)
        if chars_on_page > chars_per_page:
            page_number += 1
            chars_on_page = len(text)

        # Add the content with line and page numbers
        # plain_text.append({
        #     "line_number": line_number,
        #     "page_number": page_number,
        #     "content": text,
        #     "type": content_type
        # })
        plain_text.append(text)
        line_number += 1

    # Process document elements in order
    for element in document.element.body:
        if element.tag.endswith('p'):  # Paragraph
            # Get the paragraph text directly from the element
            text = "".join(t.text for t in element.findall('.//w:t', {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
            if text.strip():
                add_content(text, "paragraph")
            else:
                add_content("", "blank_line")

        elif element.tag.endswith('tbl'):  # Table
            table_index = len([e for e in document.element.body[:document.element.body.index(element)]
                               if e.tag.endswith('tbl')])
            table = document.tables[table_index]

            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if row_text:
                    add_content(" | ".join(row_text), "table_row")

        elif element.tag.endswith('pict') or element.tag.endswith('drawing'):  # Images
            add_content("[Image placeholder]", "image")
    # print('text1--', plain_text)
    return plain_text

def docx_to_plain_text(temp_file_path):
    from docx import Document
    import os

    # Load the document
    document = Document(temp_file_path)

    # Initialize trackers
    plain_text = []
    line_number = 1
    page_number = 1
    chars_on_page = 0
    chars_per_page = 3500  # Approximate characters per page

    # Function to add content and track page numbers
    def add_content(text, content_type="text"):
        nonlocal line_number, page_number, chars_on_page

        # Calculate page breaks based on character count
        chars_on_page += len(text)
        if chars_on_page > chars_per_page:
            page_number += 1
            chars_on_page = len(text)  # Reset for new page

            # Add a marker to indicate a new page
            # plain_text.append(f"[Page {page_number} starts here]")  # Optional marker for page start

        # Add the content with line and page numbers
        plain_text.append({
            "line_number": line_number,
            "page_number": page_number,
            "content": text,
            "type": content_type
        })
        line_number += 1

    # Process document elements in order
    for element in document.element.body:
        if element.tag.endswith('p'):  # Paragraph
            # Get the paragraph text directly from the element
            text = "".join(t.text for t in element.findall('.//w:t', {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
            if text.strip():
                add_content(text, "paragraph")
            else:
                add_content("", "blank_line")

        elif element.tag.endswith('tbl'):  # Table
            table_index = len([e for e in document.element.body[:document.element.body.index(element)]
                               if e.tag.endswith('tbl')])
            table = document.tables[table_index]

            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if row_text:
                    add_content(" | ".join(row_text), "table_row")

        elif element.tag.endswith('pict') or element.tag.endswith('drawing'):  # Images
            add_content("[Image placeholder]", "image")
    # print('text--', plain_text)
    return plain_text

# ----------
# # working without pdf
# def convert_docx_to_html(file, temp_dir, original=True):
#     doc = "org" if original else "mod"
#     def iter_block_items(parent):
#         """Iterate through the document elements (paragraphs & tables)."""
#         if isinstance(parent, _Document):
#             parent_elm = parent.element.body
#         elif isinstance(parent, _Cell):
#             parent_elm = parent._tc
#         else:
#             raise ValueError("Unexpected parent element")

#         for child in parent_elm.iterchildren():
#             if isinstance(child, CT_P):
#                 yield Paragraph(child, parent)
#             elif isinstance(child, CT_Tbl):
#                 yield Table(child, parent)

#     def render_header_subheader(block):
#         """Render headings with appropriate HTML tags inside sections."""
#         html = ""
#         if block.style.name == 'Heading 1':
#             html = f"<h1 style='font-size: 16px; margin-bottom: 5px;'>{block.text}</h1>"
#         elif block.style.name == 'Heading 2':
#             html = f"<h2 style='font-size: 15px; margin-bottom: 5px;'>{block.text}</h2>"
#         elif block.style.name == 'Heading 3':
#             html = f"<h3 style='font-size: 14px; margin-bottom: 5px;'>{block.text}</h3>"
#         elif block.style.name == 'Heading 4':
#             html = f"<h4 style='font-size: 13px; margin-bottom: 5px;'>{block.text}</h4>"
#         elif block.style.name == 'Heading 5':
#             html = f"<h5 style='font-size: 12px; margin-bottom: 5px;'>{block.text}</h5>"
#         return html

#     def render_runs(runs, outer_tag='p'):
#         """Render inline text styles inside paragraphs."""
#         html = f"<{outer_tag} style='margin-bottom: 5px; font-size: 14px;'>"
#         for run in runs:
#             text = run.text
#             if run.bold:
#                 text = f"<strong>{text}</strong>"
#             if run.italic:
#                 text = f"<em>{text}</em>"
#             if run.underline:
#                 text = f"<u>{text}</u>"
#             if hasattr(run, 'font') and run.font.size:
#                 size = run.font.size.pt
#                 text = f"<span style='font-size: {size}px;'>{text}</span>"
#             html += text
#         html += f"</{outer_tag}>"
#         return html

#     def render_list_items(items):
#         """Render list items inside an unordered list."""
#         return f"<ul style='margin-bottom: 10px;'>" + "".join(items) + "</ul>"

#     def render_table(block, document, image_path, idx):
#         """Render tables with optional row-based sections."""
#         table = block
#         # Wrap the entire table in a section
#         html = f"<section id='{doc}_section_{idx}' class='card_section'><table style='border-collapse: collapse; width: 100%; margin: 10px 0;' border='1'>"

#         for row_idx, row in enumerate(table.rows):
#             html += "<tr>"
#             for col_idx, cell in enumerate(row.cells):
#                 if row_idx == 0:  # Header row styling
#                     html += "<td style='border: 1px solid #ddd; padding: 8px; background-color: #f0f0f0; font-weight: bold;'>"
#                 else:
#                     html += "<td style='border: 1px solid #ddd; padding: 8px;'>"

#                 cbody = ""
#                 clist_items = []
#                 for cblock in iter_block_items(cell):
#                     if isinstance(cblock, Paragraph):
#                         tmp_heading_type = cblock.style.name
#                         if re.match("List\sParagraph", tmp_heading_type):
#                             clist_items.append("<li>" + cblock.text + "</li>")
#                         else:
#                             images = render_image(document, cblock, image_path)
#                             if len(clist_items) > 0:
#                                 cbody += render_list_items(clist_items)
#                                 clist_items = []
#                             if len(images) > 0:
#                                 cbody += images
#                             else:
#                                 cbody += render_runs(cblock.runs)
#                     elif isinstance(cblock, Table):
#                         cbody += render_table(cblock, document, image_path)

#                 html += cbody + "</td>"
#             html += "</tr>"

#         html += "</table></section>"
#         return html

#     def render_image(document, par, dir_path):
#         """Extract and render images inside paragraphs."""
#         ids = []
#         root = ET.fromstring(par._p.xml)
#         namespace = {
#             'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
#             'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
#             'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
#         }

#         for inline_type in ['.//wp:inline', './/wp:anchor']:
#             inlines = root.findall(inline_type, namespace)
#             for inline in inlines:
#                 imgs = inline.findall('.//a:blip', namespace)
#                 for img in imgs:
#                     id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
#                     ids.append(id)

#         response = ""
#         if ids:
#             for id in ids:
#                 image_part = document.part.related_parts[id]
#                 millis = int(round(time.time() * 1000))
#                 file_name = f"{id}-{millis}.png"
#                 file_path = os.path.join(dir_path, file_name)
#                 with open(file_path, "wb") as fr:
#                     fr.write(image_part._blob)
#                 response += f"<img style='max-width: 100%; height: auto; margin: 10px 0;' src='{file_path}'/>"
#         return response

#     document = Document(file)
#     body = ""
#     list_items = []
#     section_open = False  # Track if a section is open

#     for idx, block in enumerate(iter_block_items(document)):
#         if isinstance(block, Paragraph):
#             tmp_heading_type = block.style.name

#             # Handle empty paragraphs
#             if not block.text.strip():
#                 if not section_open:
#                     body += f'<section id="{doc}_section_{idx}" class="card_section">'
#                     section_open = True
#                 body += "<br>"
#                 continue

#             # Handle headings
#             if tmp_heading_type.startswith("Heading"):
#                 if section_open:
#                     body += "</section>"
#                 body += f'<section id="{doc}_section_{idx}" class="card_section">{render_header_subheader(block)}'
#                 section_open = True

#             # Handle list paragraphs
#             elif re.match("List\sParagraph", tmp_heading_type):
#                 list_items.append("<li>" + block.text + "</li>")

#             # Handle regular paragraphs and other content
#             else:
#                 images = render_image(document, block, temp_dir)

#                 # If we have accumulated list items, render them
#                 if list_items:
#                     if not section_open:
#                         body += f'<section id="{doc}_section_{idx}" class="card_section">'
#                         section_open = True
#                     body += render_list_items(list_items)
#                     list_items = []

#                 # Handle images
#                 if images:
#                     if not section_open:
#                         body += f'<section id="{doc}_section_{idx}" class="card_section">'
#                         section_open = True
#                     body += images

#                 # Handle regular paragraph content
#                 else:
#                     if not section_open:
#                         body += f'<section id="{doc}_section_{idx}" class="card_section">'
#                         section_open = True
#                     outer_tag = f"h{tmp_heading_type[-1]}" if 'Heading' in tmp_heading_type else 'p'
#                     body += render_runs(block.runs, outer_tag)

#         elif isinstance(block, Table):
#             if section_open:
#                 body += "</section>"
#                 section_open = False
#             body += render_table(block, document, temp_dir, idx)
#             # No need to set section_open to true here since table is self-contained in its own section

#     # Close any open section at the end
#     if section_open:
#         body += "</section>"

#     return body

# def convert_docx_to_html3(file, temp_dir, original=True):
#     """Convert DOCX to HTML with page numbers in section tags."""
#     doc = "org" if original else "mod"

#     def create_pdf(input_path):
#         """Convert DOCX to PDF using Aspose.Words and return the PDF as a binary stream."""
#         with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
#             temp_pdf_path = temp_pdf.name

#         try:
#             doc = aw.Document(input_path)
#             doc.save(temp_pdf_path)

#             with open(temp_pdf_path, "rb") as pdf_file:
#                 pdf_data = pdf_file.read()

#             return pdf_data
#         finally:
#             if os.path.exists(temp_pdf_path):
#                 os.unlink(temp_pdf_path)

#     def get_page_numbers(docx_path):
#         """Convert DOCX to PDF and get page numbers for each block of content."""
#         # Load the Word document
#         pdf_data = create_pdf(docx_path)
#         pdf_buffer = io.BytesIO(pdf_data)
#         # Open PDF with PyMuPDF
#         pdf_document = fitz.open(stream=pdf_buffer.getvalue(), filetype="pdf")

#         # Get text blocks with their page numbers
#         page_blocks = []
#         for page_num in range(len(pdf_document)):
#             page = pdf_document[page_num]
#             blocks = page.get_text("blocks")
#             for block in blocks:
#                 page_blocks.append({
#                     'text': block[4],
#                     'page': page_num + 1
#                 })

#         pdf_document.close()
#         return page_blocks

#     def find_page_number(text, page_blocks):
#         """Find the page number for a given block of text."""
#         text = text.strip()
#         if not text:
#             return 1

#         # Try to find an exact match first
#         for block in page_blocks:
#             if text in block['text']:
#                 return block['page']

#         # If no exact match, use fuzzy matching
#         best_match = None
#         highest_ratio = 0

#         for block in page_blocks:
#             ratio = difflib.SequenceMatcher(None, text, block['text']).ratio()
#             if ratio > highest_ratio:
#                 highest_ratio = ratio
#                 best_match = block

#         return best_match['page'] if best_match else 1

#     # Get page numbers for content blocks
#     page_blocks = get_page_numbers(file)

#     def iter_block_items(parent):
#         """Iterate through the document elements (paragraphs & tables)."""
#         if isinstance(parent, _Document):
#             parent_elm = parent.element.body
#         elif isinstance(parent, _Cell):
#             parent_elm = parent._tc
#         else:
#             raise ValueError("Unexpected parent element")

#         for child in parent_elm.iterchildren():
#             if isinstance(child, CT_P):
#                 yield Paragraph(child, parent)
#             elif isinstance(child, CT_Tbl):
#                 yield Table(child, parent)

#     def render_header_subheader(block):
#         """Render headings with appropriate HTML tags inside sections."""
#         html = ""
#         if block.style.name == 'Heading 1':
#             html = f"<h1 style='font-size: 16px; margin-bottom: 5px;'>{block.text}</h1>"
#         elif block.style.name == 'Heading 2':
#             html = f"<h2 style='font-size: 15px; margin-bottom: 5px;'>{block.text}</h2>"
#         elif block.style.name == 'Heading 3':
#             html = f"<h3 style='font-size: 14px; margin-bottom: 5px;'>{block.text}</h3>"
#         elif block.style.name == 'Heading 4':
#             html = f"<h4 style='font-size: 13px; margin-bottom: 5px;'>{block.text}</h4>"
#         elif block.style.name == 'Heading 5':
#             html = f"<h5 style='font-size: 12px; margin-bottom: 5px;'>{block.text}</h5>"
#         return html

#     def render_runs(runs, outer_tag='p'):
#         """Render inline text styles inside paragraphs."""
#         html = f"<{outer_tag} style='margin-bottom: 5px; font-size: 14px;'>"
#         for run in runs:
#             text = run.text
#             if run.bold:
#                 text = f"<strong>{text}</strong>"
#             if run.italic:
#                 text = f"<em>{text}</em>"
#             if run.underline:
#                 text = f"<u>{text}</u>"
#             if hasattr(run, 'font') and run.font.size:
#                 size = run.font.size.pt
#                 text = f"<span style='font-size: {size}px;'>{text}</span>"
#             html += text
#         html += f"</{outer_tag}>"
#         return html

#     def render_list_items(items):
#         """Render list items inside an unordered list."""
#         return f"<ul style='margin-bottom: 10px;'>" + "".join(items) + "</ul>"

#     def render_table(block, document, image_path, idx, page_blocks):
#         """Render tables with optional row-based sections."""
#         table = block
#         page_num = find_page_number(table.rows[0].cells[0].text if table.rows and table.rows[0].cells else '',
#                                     page_blocks)
#         html = f"<section id='{doc}_section_{idx}' class='card_section' page-number='{page_num}'>"
#         html += "<table style='border-collapse: collapse; width: 100%; margin: 10px 0;' border='1'>"
#         for row_idx, row in enumerate(table.rows):
#             html += "<tr>"
#             for col_idx, cell in enumerate(row.cells):
#                 if row_idx == 0:  # Header row styling
#                     html += "<td style='border: 1px solid #ddd; padding: 8px; background-color: #f0f0f0; font-weight: bold;'>"
#                 else:
#                     html += "<td style='border: 1px solid #ddd; padding: 8px;'>"
#                 # html += "<td style='border: 1px solid #ddd; padding: 8px;'>"
#                 cbody = ""
#                 for cblock in iter_block_items(cell):
#                     if isinstance(cblock, Paragraph):
#                         cbody += render_runs(cblock.runs)
#                     elif isinstance(cblock, Table):
#                         cbody += render_table(cblock, document, image_path, idx, page_blocks)
#                 html += cbody + "</td>"
#             html += "</tr>"
#         html += "</table></section>"
#         return html

#     def render_image(document, par, dir_path):
#         """Extract and render images inside paragraphs."""
#         ids = []
#         root = ET.fromstring(par._p.xml)
#         namespace = {
#             'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
#             'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
#             'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
#         }

#         for inline_type in ['.//wp:inline', './/wp:anchor']:
#             inlines = root.findall(inline_type, namespace)
#             for inline in inlines:
#                 imgs = inline.findall('.//a:blip', namespace)
#                 for img in imgs:
#                     id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
#                     ids.append(id)

#         response = ""
#         if ids:
#             for id in ids:
#                 image_part = document.part.related_parts[id]
#                 millis = int(round(time.time() * 1000))
#                 file_name = f"{id}-{millis}.png"
#                 file_path = os.path.join(dir_path, file_name)
#                 with open(file_path, "wb") as fr:
#                     fr.write(image_part._blob)
#                 response += f"<img style='max-width: 100%; height: auto; margin: 10px 0;' src='{file_path}'/>"
#         return response

#     document = Document(file)
#     body = ""
#     list_items = []
#     current_section = {
#         'content': [],
#         'id': None,
#         'page_num': None
#     }

#     def start_new_section(idx, page_num, content=""):
#         """Start a new section with given index and page number."""
#         nonlocal body, current_section

#         # First, close the current section if it exists
#         if current_section['content']:
#             section_content = "".join(current_section['content']).strip()
#             if section_content:
#                 body += f'<section id="{doc}_section_{current_section["id"]}" class="card_section" page-number="{current_section["page_num"]}">'
#                 body += section_content
#                 body += "</section>"

#         # Reset current section
#         current_section = {
#             'content': [content] if content else [],
#             'id': idx,
#             'page_num': page_num
#         }

#     def add_to_current_section(content):
#         """Add content to current section."""
#         current_section['content'].append(content)

#     for idx, block in enumerate(iter_block_items(document)):
#         if isinstance(block, Paragraph):
#             page_num = find_page_number(block.text, page_blocks)

#             # Handle empty paragraphs
#             if not block.text.strip():
#                 add_to_current_section("<br>")
#                 continue

#             if block.style.name.startswith("Heading"):
#                 # Start new section for heading
#                 start_new_section(idx, page_num, render_header_subheader(block))

#             elif re.match("List\sParagraph", block.style.name):
#                 list_items.append("<li>" + block.text + "</li>")

#             else:
#                 images = render_image(document, block, temp_dir)

#                 # Handle pending list items
#                 if list_items:
#                     add_to_current_section(render_list_items(list_items))
#                     list_items = []

#                 # Handle images and paragraph content
#                 if images:
#                     if not current_section['id']:
#                         start_new_section(idx, page_num)
#                     add_to_current_section(images)
#                 else:
#                     outer_tag = f"h{block.style.name[-1]}" if 'Heading' in block.style.name else 'p'
#                     content = render_runs(block.runs, outer_tag)

#                     # Start new section if none exists
#                     if not current_section['id']:
#                         start_new_section(idx, page_num)
#                     add_to_current_section(content)

#         elif isinstance(block, Table):
#             # First close any existing section
#             start_new_section(idx, find_page_number("", page_blocks))

#             # Tables are self-contained sections
#             body += render_table(block, document, temp_dir, idx, page_blocks)

#             # Reset current section after table
#             current_section = {
#                 'content': [],
#                 'id': None,
#                 'page_num': None
#             }

#     # Close final section if needed
#     if current_section['content']:
#         section_content = "".join(current_section['content']).strip()
#         if section_content:
#             body += f'<section id="{doc}_section_{current_section["id"]}" class="card_section" page-number="{current_section["page_num"]}">'
#             body += section_content
#             body += "</section>"

#     return body

# # working final
def convert_docx_to_html(file, temp_dir, original=True):
    """Convert DOCX to HTML with page numbers in section tags."""
    doc = "org" if original else "mod"
    system = platform.system()

    def create_pdf(input_path):
        """Convert DOCX to PDF using Aspose.Words and return the PDF as a binary stream."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name

        try:
            doc = aw.Document(input_path)
            doc.save(temp_pdf_path)

            with open(temp_pdf_path, "rb") as pdf_file:
                pdf_data = pdf_file.read()

            return pdf_data
        finally:
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)

    def create_pdf_in_linux(input_path):
        """Convert DOCX to PDF using LibreOffice, return PDF bytes, and delete the PDF after conversion."""
        
        output_path = os.path.splitext(input_path)[0] + ".pdf"

        try:
            # Run LibreOffice to convert DOCX to PDF
            subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", os.path.dirname(input_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
            )

            # Ensure PDF was created
            if not os.path.exists(output_path):
                raise FileNotFoundError(f"PDF conversion failed: {output_path} not found.")

            # Read PDF bytes
            with open(output_path, "rb") as pdf_file:
                pdf_data = pdf_file.read()

                return pdf_data

        except subprocess.CalledProcessError as e:
            print("Error during conversion:", e.stderr)
            raise

        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
            
    def get_page_numbers(docx_path):
        """Convert DOCX to PDF and get page numbers for each block of content."""
        # Load the Word document
        if system == "Linux":
            pdf_data = create_pdf_in_linux(docx_path)
        else:
            pdf_data = create_pdf(docx_path)

        pdf_buffer = io.BytesIO(pdf_data)
        # Open PDF with PyMuPDF
        pdf_document = fitz.open(stream=pdf_buffer.getvalue(), filetype="pdf")

        # Get text blocks with their page numbers
        page_blocks = []
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            blocks = page.get_text("blocks")
            for block in blocks:
                page_blocks.append({
                    'text': block[4],
                    'page': page_num + 1
                })

        pdf_document.close()
        return page_blocks

    def find_page_number(text, page_blocks):
        """Find the page number for a given block of text."""
        text = text.strip()
        if not text:
            return 1

        # Try to find an exact match first
        for block in page_blocks:
            if text in block['text']:
                return block['page']

        # If no exact match, use fuzzy matching
        best_match = None
        highest_ratio = 0

        for block in page_blocks:
            ratio = difflib.SequenceMatcher(None, text, block['text']).ratio()
            if ratio > highest_ratio:
                highest_ratio = ratio
                best_match = block

        return best_match['page'] if best_match else 1

    # Get page numbers for content blocks
    page_blocks = get_page_numbers(file)

    def iter_block_items(parent):
        """Iterate through the document elements (paragraphs & tables)."""
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Unexpected parent element")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def render_header_subheader(block):
        """Render headings with appropriate HTML tags inside sections."""
        html = ""
        if block.style.name == 'Heading 1':
            html = f"<h1 style='font-size: 16px; margin-bottom: 5px;'>{block.text}</h1>"
        elif block.style.name == 'Heading 2':
            html = f"<h2 style='font-size: 15px; margin-bottom: 5px;'>{block.text}</h2>"
        elif block.style.name == 'Heading 3':
            html = f"<h3 style='font-size: 14px; margin-bottom: 5px;'>{block.text}</h3>"
        elif block.style.name == 'Heading 4':
            html = f"<h4 style='font-size: 13px; margin-bottom: 5px;'>{block.text}</h4>"
        elif block.style.name == 'Heading 5':
            html = f"<h5 style='font-size: 12px; margin-bottom: 5px;'>{block.text}</h5>"
        return html

    def render_runs(runs, outer_tag='p'):
        """Render inline text styles inside paragraphs."""
        html = f"<{outer_tag} style='margin-bottom: 5px; font-size: 14px;'>"
        for run in runs:
            text = run.text
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            if hasattr(run, 'font') and run.font.size:
                size = run.font.size.pt
                text = f"<span style='font-size: {size}px;'>{text}</span>"
            html += text
        html += f"</{outer_tag}>"
        return html

    def render_list_items(items):
        """Render list items inside an unordered list."""
        return f"<ul style='margin-bottom: 10px;'>" + "".join(items) + "</ul>"

    def render_table(block, document, image_path, idx, page_blocks):
        """Render tables with optional row-based sections."""
        table = block
        page_num = find_page_number(table.rows[0].cells[0].text if table.rows and table.rows[0].cells else '',
                                    page_blocks)
        html = f"<section id='{doc}_section_{idx}' class='card_section' page-number='{page_num}'>"
        html += "<table style='border-collapse: collapse; width: 100%; margin: 10px 0;' border='1'>"
        for row_idx, row in enumerate(table.rows):
            html += "<tr>"
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:  # Header row styling
                    html += "<td style='border: 1px solid #ddd; padding: 8px; background-color: #f0f0f0; font-weight: bold;'>"
                else:
                    html += "<td style='border: 1px solid #ddd; padding: 8px;'>"
                # html += "<td style='border: 1px solid #ddd; padding: 8px;'>"
                cbody = ""
                for cblock in iter_block_items(cell):
                    if isinstance(cblock, Paragraph):
                        cbody += render_runs(cblock.runs)
                    elif isinstance(cblock, Table):
                        cbody += render_table(cblock, document, image_path, idx, page_blocks)
                html += cbody + "</td>"
            html += "</tr>"
        html += "</table></section>"
        return html

    def render_image(document, par, dir_path):
        """Extract and render images inside paragraphs."""
        ids = []
        root = ET.fromstring(par._p.xml)
        namespace = {
            'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
            'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        }

        for inline_type in ['.//wp:inline', './/wp:anchor']:
            inlines = root.findall(inline_type, namespace)
            for inline in inlines:
                imgs = inline.findall('.//a:blip', namespace)
                for img in imgs:
                    id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
                    ids.append(id)

        response = ""
        if ids:
            for id in ids:
                image_part = document.part.related_parts[id]
                millis = int(round(time.time() * 1000))
                file_name = f"{id}-{millis}.png"
                file_path = os.path.join(dir_path, file_name)
                with open(file_path, "wb") as fr:
                    fr.write(image_part._blob)
                response += f"<img style='max-width: 100%; height: auto; margin: 10px 0;' src='{file_path}'/>"
        return response

    # # Validate DOCX file
    # if not is_valid_docx(file):
    #     raise ValueError("Invalid file format. Please upload a valid DOCX file.")
    
    document = Document(file)
    body = ""
    list_items = []
    current_section = {
        'content': [],
        'id': None,
        'page_num': None
    }

    def start_new_section(idx, page_num, content=""):
        """Start a new section with given index and page number."""
        nonlocal body, current_section

        # First, close the current section if it exists
        if current_section['content']:
            section_content = "".join(current_section['content']).strip()
            if section_content:
                body += f'<section id="{doc}_section_{current_section["id"]}" class="card_section" page-number="{current_section["page_num"]}">'
                body += section_content
                body += "</section>"

        # Reset current section
        current_section = {
            'content': [content] if content else [],
            'id': idx,
            'page_num': page_num
        }

    def add_to_current_section(content):
        """Add content to current section."""
        current_section['content'].append(content)

    for idx, block in enumerate(iter_block_items(document)):
        if isinstance(block, Paragraph):
            page_num = find_page_number(block.text, page_blocks)

            # Handle empty paragraphs
            if not block.text.strip():
                add_to_current_section("<br>")
                continue

            if block.style.name.startswith("Heading"):
                # Start new section for heading
                start_new_section(idx, page_num, render_header_subheader(block))

            elif re.match("List\sParagraph", block.style.name):
                list_items.append("<li>" + block.text + "</li>")

            else:
                images = render_image(document, block, temp_dir)

                # Handle pending list items
                if list_items:
                    add_to_current_section(render_list_items(list_items))
                    list_items = []

                # Handle images and paragraph content
                if images:
                    if not current_section['id']:
                        start_new_section(idx, page_num)
                    add_to_current_section(images)
                else:
                    outer_tag = f"h{block.style.name[-1]}" if 'Heading' in block.style.name else 'p'
                    content = render_runs(block.runs, outer_tag)

                    # Start new section if none exists
                    if not current_section['id']:
                        start_new_section(idx, page_num)
                    add_to_current_section(content)

        elif isinstance(block, Table):
            # First close any existing section
            start_new_section(idx, find_page_number("", page_blocks))

            # Tables are self-contained sections
            body += render_table(block, document, temp_dir, idx, page_blocks)

            # Reset current section after table
            current_section = {
                'content': [],
                'id': None,
                'page_num': None
            }

    # Close final section if needed
    if current_section['content']:
        section_content = "".join(current_section['content']).strip()
        if section_content:
            body += f'<section id="{doc}_section_{current_section["id"]}" class="card_section" page-number="{current_section["page_num"]}">'
            body += section_content
            body += "</section>"

    return body
# ----------

def compare_documents_difflib(doc1_lines, doc2_lines):
    changes = []

    # Create lists of lines for each document
    # doc1_lines = [item['content'] for item in doc1]
    # doc2_lines = [item['content'] for item in doc2]

    # Initialize SequenceMatcher
    matcher = difflib.SequenceMatcher(None, doc1_lines, doc2_lines)

    # Process the differences
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        change = {
            'type': tag,
            'doc1_start': i1,
            'doc1_end': i2,
            'doc2_start': j1,
            'doc2_end': j2,
            'doc1_content': doc1_lines[i1:i2],
            'doc2_content': doc2_lines[j1:j2],
            'doc1_key': f"doc1_{tag}_{i1}_{i2}",
            'doc2_key': f"doc2_{tag}_{j1}_{j2}",
        }

        if tag != 'equal':  # Only add non-matching sections
            # Add more detailed information based on the type of change
            if tag == 'replace':
                change['description'] = f"Content replaced from positions {i1}-{i2} in doc1 to {j1}-{j2} in doc2"
            elif tag == 'delete':
                change['description'] = f"Content deleted from positions {i1}-{i2} in doc1"
            elif tag == 'insert':
                change['description'] = f"Content inserted at positions {j1}-{j2} in doc2"

            changes.append(change)

    return changes

# ----------

def diff_html(html1, html2, differences):
    # print('differences----', differences)
    highlight_html1 = html1
    highlight_html2 = html2

    STYLE_MAP = {
        "replace": "rgb(140, 190, 240)",
        "insert": "rgb(132, 232, 159)",
        "delete": "rgb(255, 180, 180)",
    }

    TABLE_CELL_TAG = "</p></td><td style='border: 1px solid #ddd; padding: 8px;'><p style='margin-bottom: 5px; font-size: 14px;'>"

    def highlight_text(id, content, color_style):
        # key = next((k for k, v in STYLE_MAP.items() if v == color_style), None)
        if '|' in content:
            table_row = content.replace(' | ', TABLE_CELL_TAG)
            table_row_list = table_row.split(TABLE_CELL_TAG)
            return TABLE_CELL_TAG.join(
                [f'<span id="{id}" style="background: {color_style};">{word}</span>' for word in table_row_list])
        return f'<span id="{id}" style="background: {color_style};">{content}</span>'

    for id, diff in enumerate(differences):
        color_style = STYLE_MAP.get(diff['type'])
        if not color_style:
            continue
        highlighted1 = ''
        highlighted2 = ''
        for content in diff['doc1_content']:
            id = diff['doc1_key']
            highlighted = highlight_text(id, content, color_style)
            # highlighted1 = highlighted
            # print("HTML1--", highlight_html1) if id == 'doc1_replace_81_82' else None
            # print("REPLACE THIS--", content.replace(' | ', TABLE_CELL_TAG)) if highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)) < 1 else None
            # print("CONTENT--", content) if highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)) < 1 else None
            # print("INDEX--", highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)))
            # print(highlight_html1) if highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)) < 1 else None
            highlight_html1 = highlight_html1.replace(content.replace(' | ', TABLE_CELL_TAG), highlighted, 1)
            # print("WITH THIS--", highlighted) if id == 'doc1_replace_81_82' else None

        for content in diff['doc2_content']:
            id = diff['doc2_key']
            # highlighted = highlight_text(id, content, color_style if diff['type'] != 'delete' else STYLE_MAP['replace'])
            highlighted = highlight_text(id, content, color_style)
            # highlighted2 = highlighted
            highlight_html2 = highlight_html2.replace(content.replace(' | ', TABLE_CELL_TAG), highlighted, 1)
        # print('1>>', id, highlighted1)
        # print('2>>', id, highlighted2)

    html_text1 = highlight_html1
        # f"""
        # <div style="
        #     background-color: #fff;
        #     margin: auto;
        #     overflow-y: auto;
        # ">
        #     {highlight_html1}
        # </div>
        # """
    html_text2 = highlight_html2
        # f"""
        # <div style="
        #     background-color: #fff;
        #     margin: auto;
        #     overflow-y: auto;
        # ">
        #     {highlight_html2}
        # </div>
        # """

    return {'html_text1': html_text1, 'html_text2': html_text2}

"""
width: 794px; 
padding: 20px; 
border: 1px solid #ddd; 
border-radius: 5px;
box-shadow: 0px 4px 6px rgba(0,0,0,0.1);
height: 600px;
"""

# ----------

def format_html_to_lines_bs4(html_text):
    soup = BeautifulSoup(html_text, "html.parser")
    formatted_lines = []

    def traverse(element):
        """Recursively process HTML elements while preserving attributes"""
        if isinstance(element, NavigableString):  # If it's text, add it directly
            text = element.strip()
            if text:  # Avoid empty lines
                formatted_lines.append(text)
            return

        if element.name:  # If it's a tag
            # Extract attributes (including style)
            attrs = " ".join([f'{key}="{value}"' for key, value in element.attrs.items()])
            tag_open = f"<{element.name}{' ' + attrs if attrs else ''}>"
            formatted_lines.append(tag_open)  # Open tag

        for child in element.children:  # Traverse child elements
            traverse(child)

        if element.name and element.name not in ["br", "img", "input", "meta", "hr", "link"]:  # Avoid closing self-closing tags
            formatted_lines.append(f"</{element.name}>")

    for child in soup.contents:  # Start processing from root elements
        traverse(child)

    return formatted_lines

# ----------

def highlight_html(html1, html2, differences):
    # print('differences----', differences)
    html_lines1 = format_html_to_lines_bs4(html1)
    html_lines2 = format_html_to_lines_bs4(html2)
    # return {}

    STYLE_MAP = {
        "replace": "rgb(189, 224, 242)",
        "insert": "rgb(133, 242, 166)",
        "delete": "rgb(251, 214, 212)",
    }
    #
    # TABLE_CELL_TAG = "</p></td><td style='border: 1px solid #ddd; padding: 8px;'><p style='margin-bottom: 5px; font-size: 14px;'>"
    #
    # def highlight_text(id, content, color_style):
    #     # key = next((k for k, v in STYLE_MAP.items() if v == color_style), None)
    #     if '|' in content:
    #         table_row = content.replace(' | ', TABLE_CELL_TAG)
    #         table_row_list = table_row.split(TABLE_CELL_TAG)
    #         return TABLE_CELL_TAG.join(
    #             [f'<span id="{id}" style="background: {color_style};">{word}</span>' for word in table_row_list])
    #     return f'<span id="{id}" style="background: {color_style};">{content}</span>'
    #
    for diff in reversed(differences):
        color_style = STYLE_MAP.get(diff['type'])
        if not color_style:
            continue
        id = diff['doc1_key']
        html_list = []
        for line in html_lines1[diff['doc1_start']:diff['doc1_end']]:
            if '<' in line and '>' in line:
                html_list.append(line)
            else:
                html_list.append(f'<span id="{id}" style="background: {color_style};">{line}</span>')
        html_lines1[diff['doc1_start']:diff['doc1_end']] = html_list
        # print('1---', html_lines1)

        id = diff['doc2_key']
        html_list = []
        for line in html_lines2[diff['doc2_start']:diff['doc2_end']]:
            if '<' in line and '>' in line:
                html_list.append(line)
            else:
                html_list.append(f'<span id="{id}" style="background: {color_style};">{line}</span>')
        html_lines2[diff['doc2_start']:diff['doc2_end']] = html_list
        # print('2---', html_lines2)

    html_highlight1 = ''.join(html_lines1)
    html_highlight2 = ''.join(html_lines2)

    return {'html_text1': html_highlight1, 'html_text2': html_highlight2}
            # highlighted = highlight_text(id, content, color_style)
            # highlighted1 = highlighted
    #         # print("HTML1--", highlight_html1) if id == 'doc1_replace_81_82' else None
    #         # print("REPLACE THIS--", content.replace(' | ', TABLE_CELL_TAG)) if highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)) < 1 else None
    #         # print("CONTENT--", content) if highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)) < 1 else None
    #         print("INDEX--", highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)))
    #         # print(highlight_html1) if highlight_html1.find(content.replace(' | ', TABLE_CELL_TAG)) < 1 else None
    #         highlight_html1 = highlight_html1.replace(content.replace(' | ', TABLE_CELL_TAG), highlighted, 1)
    #         # print("WITH THIS--", highlighted) if id == 'doc1_replace_81_82' else None
    #
    #     for content in diff['doc2_content']:
    #         id = diff['doc2_key']
    #         # highlighted = highlight_text(id, content, color_style if diff['type'] != 'delete' else STYLE_MAP['replace'])
    #         highlighted = highlight_text(id, content, color_style)
    #         # highlighted2 = highlighted
    #         highlight_html2 = highlight_html2.replace(content.replace(' | ', TABLE_CELL_TAG), highlighted, 1)
    #     # print('1>>', id, highlighted1)
    #     # print('2>>', id, highlighted2)
    #
    # html_text1 = \
    #     f"""
    #     <div style="
    #         background-color: #fff;
    #         margin: auto;
    #         overflow-y: auto;
    #     ">
    #         {highlight_html1}
    #     </div>
    #     """
    # html_text2 = \
    #     f"""
    #     <div style="
    #         background-color: #fff;
    #         margin: auto;
    #         overflow-y: auto;
    #     ">
    #         {highlight_html2}
    #     </div>
    #     """
    #
    # return {'html_text1': html_text1, 'html_text2': html_text2}

def find_html_match(original_text, modified_text):
    result = []

    # Track already matched items
    matched_in_list2 = set()
    matched_in_list1 = set()  # To track items in list1 that have been processed

    # Initialize the serial number
    id_no = 1

    for idx1, item1 in enumerate(original_text):
        item1_content = item1['content'].replace('\n', ' ').strip()  # Remove any newline characters for comparison
        matched = False

        for idx2, item2 in enumerate(modified_text):
            item2_content = item2['content'].replace('\n', ' ').strip()  # Remove any newline characters for comparison

            if item1_content == item2_content and idx2 not in matched_in_list2:  # Exact match
                result.append({
                    'id_no': id_no,
                    'type': 'exact_matched',
                    'original': item1_content,
                    'modified': item2_content,
                    'org_idx': item1['line_no'],
                    'mod_idx': item2['line_no']
                })
                matched_in_list2.add(idx2)  # Mark this element as matched
                matched = True
                id_no += 1  # Increment serial number
                break
            elif item1_content in item2_content or item2_content in item1_content and idx2 not in matched_in_list2:  # Partial match (replaced)
                result.append({
                    'id_no': id_no,
                    'type': 'replaced',
                    'original': item1_content,
                    'modified': item2_content,
                    'org_idx': item1['line_no'],
                    'mod_idx': item2['line_no']
                })
                matched_in_list2.add(idx2)
                matched = True
                id_no += 1  # Increment serial number
                break

        # If no match found, no need to append anything for this item in list1
        if not matched:
            result.append({
                'id_no': id_no,
                'type': 'removed',
                'original': item1_content,
                'modified': '',
                'org_idx': item1['line_no'],
                'mod_idx': -1  # No corresponding index in list2
            })
            id_no += 1  # Increment serial number

    # Check for "Added" elements in list2 that aren't in list1
    for idx2, item2 in enumerate(modified_text):
        if idx2 not in matched_in_list2:
            result.append({
                'id_no': id_no,
                'type': 'added',
                'original': '',
                'modified': item2['content'],
                'org_idx': -1,  # No original index because it's a new addition
                'mod_idx': item2['line_no']
            })
            id_no += 1  # Increment serial number

    return result  # Return the result list instead of an empty dictionary

def highlight_all_html(original_text, modified_text, check_list):
    for check in reversed(check_list):
        check_type = check['type']
        check_id = check['id_no']
        content1 = check.get('original', '')
        content2 = check.get('modified', '')
        if check_type in {'exact_matched', 'replaced'}:
            org_id = f"original_{check_type}_{check_id}"
            mod_id = f"modified_{check_type}_{check_id}"
            # tag1 = f"<span id='{org_id}' class='{check_type}_highlight_span'><a class='highlight_a' href='#{mod_id}'>{content1}</a></span>"
            # tag2 = f"<span id='{mod_id}' class='{check_type}_highlight_span'><a class='highlight_a' href='#{org_id}'>{content2}</a></span>"
            tag1 = f"<span id='{org_id}' class='{check_type}_highlight_span'>{content1}</span>"
            tag2 = f"<span id='{mod_id}' class='{check_type}_highlight_span'>{content2}</span>"

            original_text[check['org_idx'] - 1] = tag1
            modified_text[check['mod_idx'] - 1] = tag2

        elif check_type == 'added':
            mod_id = f"modified_{check_type}_{check_id}"
            # tag2 = f"<span id='{mod_id}' class='{check_type}_highlight_span'><a class='highlight_a' href='#'>{content2}</a></span>"
            tag2 = f"<span id='{mod_id}' class='{check_type}_highlight_span'>{content2}</span>"
            modified_text[check['mod_idx'] - 1] = tag2

        elif check_type == 'removed':
            org_id = f"original_{check_type}_{check_id}"
            # tag1 = f"<span id='{org_id}' class='{check_type}_highlight_span'><a class='highlight_a' href='#'>{content1}</a></span>"
            tag1 = f"<span id='{org_id}' class='{check_type}_highlight_span'>{content1}</span>"
            original_text[check['org_idx'] - 1] = tag1

    return {"original_text": original_text, "modified_text": modified_text}

# ----------

def extract_sections_from_html(data):
    soup = BeautifulSoup(data, 'html.parser')

    sections = []
    for section in soup.find_all('section'):
        section_id = section.get('id')
        content = ' '.join(section.stripped_strings)
        sections.append({section_id: content})

    return sections

# ----------

def find_section_in_section_html(html_content, section_id):
    soup = BeautifulSoup(html_content, "html.parser")
    section = soup.find("section", {"id": f"{section_id}"})

    if section:
        html_text = str(section)  # Converts the section element back to an HTML string
        return html_text
    else:
        print("Section not found")

# ----------

def find_and_replace_section_in_section_html(section_html, section_id, styled_section_html):
    # Parse the HTML
    soup = BeautifulSoup(section_html, "html.parser")

    # Find the section with the specified ID
    target_section = soup.find("section", {"id": section_id})
    html_string = '''class="['card_section', 'replaced']"'''
    replace_string = '''class="card_section replaced"'''
    styled_section_html = styled_section_html.replace(html_string, replace_string)

    # If found, replace it with new content
    if target_section:
        new_section_soup = BeautifulSoup(styled_section_html, "html.parser")
        target_section.replace_with(new_section_soup)
        return str(soup)
    else:
        print("Section not found")

# ----------

def remove_empty_sections(html):
    # Parse the HTML content
    soup = BeautifulSoup(html, 'html.parser')

    # Find all <section> tags
    sections = soup.find_all('section')

    # Loop through each <section> tag
    for section in sections:
        # Check if the section is empty or contains only whitespace and <br> tags
        if not section.get_text(strip=True) and not section.find(lambda tag: tag.name not in ['br']):
            section.decompose()  # Remove the empty section
        # Check if the section contains only an empty table
        elif section.find('table') and not section.find('table').find_all(['tr', 'td', 'th']):
            section.decompose()  # Remove section if it only contains an empty <table>

    return str(soup)


def analyze_section_changes(original_lines, modified_lines, table_content=False, same_length=False, model=None, similarity_threshold=0.999999):
    """Compare two document versions and identify matched, removed, and added sections."""

    if model is None:
        model = SentenceTransformer('all-MiniLM-L6-v2')

    # Handle empty inputs
    if not original_lines and not modified_lines:
        return []  # Nothing to compare

    if not original_lines:
        return [{'original_content': '', 'modified_content': mod, 'org_idx': -1, 'mod_idx': idx, 'type': 'added'}
                for idx, mod in enumerate(modified_lines)]

    if not modified_lines:
        return [{'original_content': orig, 'modified_content': '', 'org_idx': idx, 'mod_idx': -1, 'type': 'removed'}
                for idx, orig in enumerate(original_lines)]

    original_data = [{"index": idx, "content": content} for idx, content in enumerate(original_lines)]
    modified_data = [{"index": idx, "content": content} for idx, content in enumerate(modified_lines)]

    original_contents = [item['content'] for item in original_data]
    modified_contents = [item['content'] for item in modified_data]

    # Compute embeddings
    original_embeddings = model.encode(original_contents, convert_to_tensor=True)
    modified_embeddings = model.encode(modified_contents, convert_to_tensor=True)

    # Compute similarity matrix
    similarity_matrix = torch.nn.functional.cosine_similarity(
        original_embeddings.unsqueeze(1),
        modified_embeddings.unsqueeze(0),
        dim=2
    )

    results = []

    if table_content and not same_length:
        matched_originals = set()  # Track matched original indices
        matched_modifieds = set()  # Track matched modified indices

        # Process original content to detect matched and removed sections
        for i, orig_item in enumerate(original_data):
            similarities = similarity_matrix[i]
            best_match_idx = torch.argmax(similarities).item()
            max_similarity = similarities[best_match_idx].item()

            if max_similarity >= similarity_threshold and best_match_idx not in matched_modifieds:
                results.append({
                    'original_content': orig_item['content'],
                    'modified_content': modified_data[best_match_idx]['content'],
                    'org_idx': orig_item['index'],
                    'mod_idx': best_match_idx,
                    'type': 'matched'
                })
                matched_originals.add(orig_item['index'])
                matched_modifieds.add(best_match_idx)
            else:
                results.append({
                    'original_content': orig_item['content'],
                    'modified_content': '',
                    'org_idx': orig_item['index'],
                    'mod_idx': -1,
                    'type': 'removed'
                })

        # Process modified content to detect added sections
        for j, mod_item in enumerate(modified_data):
            if j not in matched_modifieds:
                results.append({
                    'original_content': '',
                    'modified_content': mod_item['content'],
                    'org_idx': -1,
                    'mod_idx': mod_item['index'],
                    'type': 'added'
                })

    else:
        for i, orig_item in enumerate(original_data):
            similarities = similarity_matrix[i]
            best_match_idx = torch.argmax(similarities).item()
            max_similarity = similarities[best_match_idx].item()

            result_type = 'matched' if max_similarity >= similarity_threshold else 'removed'

            results.append({
                'original_content': orig_item['content'],
                'modified_content': modified_data[best_match_idx]['content'],
                'org_idx': i,
                'mod_idx': best_match_idx,
                'type': result_type
            })

    return results

def group_by_sections(data, original_section_style, modified_section_style):
    grouped_data = []
    current_section = []
    last_type = None

    # Step 1: Group data sequentially while preserving order
    for entry in data:
        if last_type is None or entry["type"] == last_type:
            current_section.append(entry)
        else:
            grouped_data.append(current_section)
            current_section = [entry]
        last_type = entry["type"]

    if current_section:
        grouped_data.append(current_section)

    # Step 2: Process grouped data into sectioned groups
    sectioned_groups = {}
    for i, section in enumerate(grouped_data):
        data_category = "matched" if any(entry["type"] == "matched" for entry in section) else "replaced"

        style1 = original_section_style.format(i).replace(
            'class="card_section replaced" data-category="replaced"',
            f'class="card_section {data_category}" data-category="{data_category}"'
        )
        style2 = modified_section_style.format(i).replace(
            'class="card_section replaced" data-category="replaced"',
            f'class="card_section {data_category}" data-category="{data_category}"'
        )

        # Step 3: Add start and end markers
        section_start = {
            "original_content": style1,
            "modified_content": style2,
            "org_idx": -1,
            "mod_idx": -1,
            "type": "start"
        }
        section_end = {
            "original_content": "</section>",
            "modified_content": "</section>",
            "org_idx": -1,
            "mod_idx": -1,
            "type": "end"
        }

        sectioned_groups[f"section_{i}"] = [section_start] + section + [section_end]

    return sectioned_groups

def generate_html_from_grouped_sections(grouped_sections):
    original_list = []
    modified_list = []

    for section in grouped_sections.values():
        for item in section:
            original_list.append(item["original_content"])
            modified_list.append(item["modified_content"])

    original_html = '\n'.join(original_list)
    modified_html = '\n'.join(modified_list)

    # print("Original HTML:\n", original_html)
    # print("Modified HTML:\n", modified_html)

    return {"original_html": original_html, "modified_html": modified_html}

def handle_table_html_using_bs4(html_text):

    soup = BeautifulSoup(html_text, "html.parser")

    for section in soup.find_all("section"):
        table = soup.new_tag("table", border="1", style="border-collapse: collapse; width: 100%; margin: 10px 0;")
        tbody = soup.new_tag("tbody")
        table.insert(0, tbody)

        for tr in section.find_all("tr"):
            tbody.append(tr.extract())

        section.insert(1, table)

    desired_output = str(soup)
    return desired_output

def is_table_content(html_text1, html_text2):
    if ('<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">' in html_text1 and
            '<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">' in html_text2):
        return True
    else:
        return False

def is_same_length_of_lines(html_line1, html_line2):
    if len(html_line1) == len(html_line2):
        return True
    else:
        return False
    
def find_section_and_section_style_in_section_html(html_content, section_id):
    soup = BeautifulSoup(html_content, "html.parser")
    section = soup.find("section", {"id": f"{section_id}"})

    if section:
        section_style = str(section).split(">", 1)[0] + ">"  # Extracts the opening <section> tag with attributes
        section_soup = BeautifulSoup(str(section), "html.parser")
        updated_html = "".join(str(tag) for tag in section_soup.section.contents)

        return {"html_text": str(section), "section_style": section_style, "updated_html": updated_html}
    else:
        print("Section not found", section_id)
        return None

def section_lines_using_bs4(html_text):
    soup = BeautifulSoup(html_text, 'html.parser')

    if '<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">' in html_text:
        # Extract all <tr> elements and convert them into a list of strings
        table_rows = [str(row) for row in soup.find_all("tr")]

        # Print the final list
        # print(table_rows)
        return table_rows

    else:
        # Extract elements and convert them into a list of strings
        elements_list = [str(element) for element in soup.find_all(["h2", "h3", "p"])]

        # Print the final list
        # print(elements_list)
        return elements_list
    
def find_and_replace_section_element(section_html, updated_html, section_id):
    soup = BeautifulSoup(section_html, "html.parser")

    # Find the section with the given id
    section = soup.find("section", {"id": section_id})

    if section:
        # Replace the found section with the updated HTML
        updated_soup = BeautifulSoup(updated_html, "html.parser")
        section.replace_with(updated_soup)
    return str(soup)

def section_lines_with_section_using_bs4(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    sections = [str(section) for section in soup.find_all("section")]
    return sections

def split_html_sections_in_list(html_original_lines):
    html_data1 = []

    for html_content in html_original_lines:
        if '<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">' in html_content:
            html_data1.append(html_content)
            continue

        soup = BeautifulSoup(html_content, "html.parser")
        original_section = soup.find("section")
        section_attributes = {attr: original_section[attr] for attr in
                              original_section.attrs} if original_section else {}

        base_id = section_attributes.get("id", "section") + "_{}"
        section_attributes.pop("id", None)  # Remove original ID to avoid duplicates

        attributes_str = " ".join(f'{key}="{value}"' for key, value in section_attributes.items())

        content_without_section = "".join(
            str(tag) for tag in original_section.contents) if original_section else html_content

        soup = BeautifulSoup(content_without_section, "html.parser")
        sections = []
        current_section = []
        section_id = 1
        consecutive_br_count = 0

        for tag in soup.contents:
            if tag.name == "br":
                consecutive_br_count += 1
            else:
                if consecutive_br_count >= 2:
                    if current_section:
                        section_html = f'<section id="{base_id.format(section_id)}" {attributes_str}>{"".join(current_section)}</section>'
                        sections.append(section_html)
                    current_section = []
                    section_id += 1
                consecutive_br_count = 0
                current_section.append(str(tag))

        if current_section:
            section_html = f'<section id="{base_id.format(section_id)}" {attributes_str}>{"".join(current_section)}</section>'
            sections.append(section_html)

        html_data1.extend(sections)

    return html_data1


def split_html_after_p_tags(html_list):

    def split_long_p_tags(groups):
        updated_groups = []
        for group in groups:
            updated_group = []
            for p_tag in group:
                soup = BeautifulSoup(str(p_tag), 'html.parser')
                p = soup.find('p')
                if p and len(p.get_text()) > 500:
                    sentences = p.get_text().split('. ')
                    for sentence in sentences:
                        if sentence.strip():
                            new_p = soup.new_tag('p', **p.attrs)
                            span = p.find('span')
                            if span:
                                span_tag = soup.new_tag('span', **span.attrs)
                                span_tag.string = sentence.strip() + ('.' if not sentence.endswith('.') else '')
                                new_p.append(span_tag)
                            else:
                                new_p.string = sentence.strip() + ('.' if not sentence.endswith('.') else '')
                            updated_group.append(new_p)
                else:
                    updated_group.append(p_tag)
            updated_groups.append(updated_group)
        return updated_groups

    updated_html_data = []

    for html_content in html_list:
        if '<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">' in html_content:
            updated_html_data.append(html_content)
            continue

        soup = BeautifulSoup(html_content, "html.parser")
        section = soup.find("section")
        if not section:
            updated_html_data.append(html_content)
            continue

        section_id = section.get("id", "")
        section_class = " ".join(section.get("class", []))
        page_number = section.get("page-number", "")

        # Handling <li> tags
        li_tags = section.find_all("li")
        if len(li_tags) >= 2:
            groups = [li_tags[i:i + 1] for i in range(0, len(li_tags), 1)]
            for i, group in enumerate(groups):
                new_section_id = f"{section_id}_li_{i + 1}"
                new_section = BeautifulSoup("<section></section>", "html.parser").section
                new_section["id"] = new_section_id
                if section_class:
                    new_section["class"] = section_class
                if page_number:
                    new_section["page-number"] = page_number

                ul = soup.new_tag("ul")
                for li in group:
                    ul.append(li)
                new_section.append(ul)

                updated_html_data.append(str(new_section))
            continue  # Skip further processing since we handled <li> separately

        # Handling <p> tags
        p_tags = section.find_all("p")
        if len(p_tags) >= 2:
            groups = [p_tags[i:i + 1] for i in range(0, len(p_tags), 1)]
            updated_groups = split_long_p_tags(groups)
            for i, group in enumerate(updated_groups):
                new_section_id = f"{section_id}_p_{i + 1}"
                new_section = BeautifulSoup("<section></section>", "html.parser").section
                new_section["id"] = new_section_id
                if section_class:
                    new_section["class"] = section_class
                if page_number:
                    new_section["page-number"] = page_number

                new_section.extend(group)
                updated_html_data.append(str(new_section))
        else:
            updated_html_data.append(html_content)

    return updated_html_data

def find_duplicate_matches_in_matched_sections(matches):
    if not matches:
        return []

    # Dictionary to count occurrences of each original section ID
    original_count = {}

    # Dictionary to count occurrences of each modified section ID
    modified_count = {}

    # First pass: Count occurrences
    for match in matches:
        orig_id = match['original_section_id']
        mod_id = match['modified_section_id']

        if orig_id not in original_count:
            original_count[orig_id] = 1
        else:
            original_count[orig_id] += 1

        if mod_id not in modified_count:
            modified_count[mod_id] = 1
        else:
            modified_count[mod_id] += 1

    # Find duplicates (sections with count > 1)
    duplicate_originals = [orig_id for orig_id, count in original_count.items() if count > 1]
    duplicate_modifieds = [mod_id for mod_id, count in modified_count.items() if count > 1]

    # Find all matches involving duplicates
    duplicate_matches = []
    for match in matches:
        if match['original_section_id'] in duplicate_originals or match['modified_section_id'] in duplicate_modifieds:
            duplicate_matches.append(match)

    return duplicate_matches

def group_duplicate_matches_in_matched_sections(duplicate_matches):
    if not duplicate_matches:
        return {}

    # Dictionary to track all matches by original section ID
    original_groups = {}

    # Dictionary to track all matches by modified section ID
    modified_groups = {}

    # First pass: Count occurrences to identify duplicates
    original_count = {}
    modified_count = {}

    for match in duplicate_matches:
        orig_id = match['original_section_id']
        mod_id = match['modified_section_id']

        if orig_id not in original_count:
            original_count[orig_id] = 1
        else:
            original_count[orig_id] += 1

        if mod_id not in modified_count:
            modified_count[mod_id] = 1
        else:
            modified_count[mod_id] += 1

    # Group matches by their IDs
    for match in duplicate_matches:
        orig_id = match['original_section_id']
        mod_id = match['modified_section_id']

        # Group by original ID if it appears multiple times
        if original_count.get(orig_id, 0) > 1:
            if orig_id not in original_groups:
                original_groups[orig_id] = []
            original_groups[orig_id].append(match)

        # Group by modified ID if it appears multiple times
        if modified_count.get(mod_id, 0) > 1:
            if mod_id not in modified_groups:
                modified_groups[mod_id] = []
            modified_groups[mod_id].append(match)

    # Combine the results
    result = {}
    result.update(original_groups)
    result.update(modified_groups)

    return result

def extract_scores_from_groups_in_matched_sections(grouped_matches):
    if not grouped_matches:
        return []

    result = {}

    # Iterate through each group
    for group_id, matches in grouped_matches.items():
        # Extract scores from matches in this group
        scores = [match['score'] for match in matches]
        result[group_id] = scores

    return result

def determine_best_matched_section(scores):
    num_students = len(scores)
    subject_winners = [0] * num_students  # Track how many subjects each student wins
    subjects = list(scores[0].keys())  # Get the subjects from the first student's dictionary

    # Compare each subject across all students
    for subject in subjects:
        subject_scores = [student[subject] for student in scores]
        max_score = max(subject_scores)

        # Award the subject win to students who have the highest score
        for i in range(num_students):
            if scores[i][subject] == max_score:
                subject_winners[i] += 1

    # Determine the student who wins in at least 2 subjects
    for i in range(num_students):
        if subject_winners[i] >= 2:
            # return f"Student {i + 1} passes with scores: {scores[i]}"
            return scores[i]

    # If there's no clear winner, check the 'length' score
    highest_length_score = max([student['length'] for student in scores])
    for i in range(num_students):
        if scores[i]['length'] == highest_length_score:
            # return f"Student {i + 1} passes based on highest 'length' score with scores: {scores[i]}"
            return scores[i]

    return "No clear winner"

def extract_plain_text_from_html(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    if '<tr>' in html_content:
        table = soup.find('table')
        rows = []
        for row in table.find_all('tr'):
            row_texts = []
            for cell in row.find_all(['td', 'th']):
                # Get text and strip whitespace
                row_texts.append(cell.get_text().strip())
            # Join the cells with pipes
            rows.append(" | ".join(row_texts))
        # Join rows with newlines
        return "\n".join(rows)
    else:
        soup = BeautifulSoup(html_content, 'html.parser')
        # Extract text from the paragraph
        paragraph_text = soup.get_text().strip()
        return paragraph_text

def remove_duplicate_sections_ids_from_filtered_matched_section(data):
    # Dictionary to store the best match based on length and entailment score
    filtered_data = {}

    # Iterate through the data
    for entry in data:
        original_id = entry["original_section_id"]

        # If the original_id is not in the dictionary or has a lower length and entailment score, update it
        if (original_id not in filtered_data or
                (entry["score"]["length"] > filtered_data[original_id]["score"]["length"] and
                 entry["score"]["entailment"] > filtered_data[original_id]["score"]["entailment"])):
            filtered_data[original_id] = entry

    # Convert the filtered dictionary back to a list
    output_data = list(filtered_data.values())
    return output_data

def merge_section_details(section_details, matched_section):
    # Create a dictionary to keep track of the best match for each original_section_id
    original_to_modified_map = {}

    # Track all original and modified IDs used in matched relationships
    matched_original_ids = set()
    matched_modified_ids = set()

    # First, process the matched items from both lists
    # Start with section_details
    for item in section_details:
        if item["change_type"] == "matched":
            org_id = item["original_section_id"]
            mod_id = item["modified_section_id"]
            original_to_modified_map[org_id] = mod_id
            matched_original_ids.add(org_id)
            matched_modified_ids.add(mod_id)

    # Then process matched_section with priority
    for item in matched_section:
        if item["change_type"] == "matched":
            org_id = item["original_section_id"]
            mod_id = item["modified_section_id"]

            # Always update with the mapping from matched_section
            original_to_modified_map[org_id] = mod_id
            matched_original_ids.add(org_id)
            matched_modified_ids.add(mod_id)

    # Track unique removed and added IDs
    unique_removed_original_ids = set()
    unique_added_modified_ids = set()

    # Combine the two lists
    combined_list = []

    # Process matched items first
    for org_id, mod_id in original_to_modified_map.items():
        combined_list.append({
            "change_type": "matched",
            "original_section_id": org_id,
            "modified_section_id": mod_id
        })

    # Process removed items from section_details
    for item in section_details:
        if item["change_type"] == "removed":
            org_id = item["original_section_id"]

            # Skip if this original ID is already in a matched relationship
            if org_id in matched_original_ids:
                continue

            # Skip if we've already processed this original ID
            if org_id in unique_removed_original_ids:
                continue

            unique_removed_original_ids.add(org_id)
            combined_list.append(item)

    # Process removed items from matched_section (with priority)
    for item in matched_section:
        if item["change_type"] == "removed":
            org_id = item["original_section_id"]

            # Skip if this original ID is already in a matched relationship
            if org_id in matched_original_ids:
                continue

            # If already in our result, update it (since matched_section has priority)
            if org_id in unique_removed_original_ids:
                # Find and update the existing item
                for existing in combined_list:
                    if (existing["change_type"] == "removed" and
                            existing["original_section_id"] == org_id):
                        existing["modified_section_id"] = item["modified_section_id"]
                        break
            else:
                # Add as new
                unique_removed_original_ids.add(org_id)
                combined_list.append(item)

    # Process added items from section_details
    for item in section_details:
        if item["change_type"] == "added":
            mod_id = item["modified_section_id"]

            # Skip if this modified ID is already in a matched relationship
            if mod_id in matched_modified_ids:
                continue

            # Skip if we've already processed this modified ID
            if mod_id in unique_added_modified_ids:
                continue

            unique_added_modified_ids.add(mod_id)
            combined_list.append(item)

    # Process added items from matched_section (with priority)
    for item in matched_section:
        if item["change_type"] == "added":
            mod_id = item["modified_section_id"]

            # Skip if this modified ID is already in a matched relationship
            if mod_id in matched_modified_ids:
                continue

            # If already in our result, update it (since matched_section has priority)
            if mod_id in unique_added_modified_ids:
                # Find and update the existing item
                for existing in combined_list:
                    if (existing["change_type"] == "added" and
                            existing["modified_section_id"] == mod_id):
                        existing["original_section_id"] = item["original_section_id"]
                        break
            else:
                # Add as new
                unique_added_modified_ids.add(mod_id)
                combined_list.append(item)

    return combined_list

# ----------

# path  = './uploads/brd_one.docx'
# a = docx_to_plain_text1(path)
# b = docx_to_plain_text(path)

